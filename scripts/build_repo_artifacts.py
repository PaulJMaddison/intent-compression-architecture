from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DIAGRAM_DIR = ROOT / "diagrams"
DOCX_PATH = ROOT / "ICA_Engineering_Design_Proposal1.docx"


COLORS = {
    "bg": "#F7F9FC",
    "panel": "#FFFFFF",
    "ink": "#1F2937",
    "muted": "#5B6470",
    "line": "#D5DDE7",
    "blue_fill": "#E8F1FB",
    "blue_stroke": "#2D5B91",
    "amber_fill": "#FFF2D9",
    "amber_stroke": "#A86400",
    "purple_fill": "#EFE8FB",
    "purple_stroke": "#6650A4",
    "green_fill": "#E9F6EE",
    "green_stroke": "#2F6F44",
    "red_fill": "#FCEBEB",
    "red_stroke": "#A12B2B",
    "teal_fill": "#E8F6F7",
    "teal_stroke": "#236E78",
}


def load_font(name: str, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts") / name.lower(),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


FONT_TITLE = load_font("arialbd.ttf", 44)
FONT_SUBTITLE = load_font("arial.ttf", 24)
FONT_BOX = load_font("arialbd.ttf", 26)
FONT_BODY = load_font("arial.ttf", 22)
FONT_SMALL = load_font("arial.ttf", 18)


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def add_shadow(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], radius: int = 22) -> None:
    x0, y0, x1, y1 = box
    shadow = (x0 + 8, y0 + 10, x1 + 8, y1 + 10)
    draw.rounded_rectangle(
        shadow,
        radius=radius,
        fill=(30, 41, 59, 24),
        outline=None,
    )


def draw_block(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    fill_key: str,
    stroke_key: str,
) -> None:
    add_shadow(draw, box)
    draw.rounded_rectangle(
        box,
        radius=22,
        fill=hex_to_rgb(COLORS[fill_key]),
        outline=hex_to_rgb(COLORS[stroke_key]),
        width=4,
    )
    x0, y0, x1, y1 = box
    center_x = (x0 + x1) / 2
    draw.text(
        (center_x, y0 + 22),
        title,
        anchor="ma",
        fill=hex_to_rgb(COLORS["ink"]),
        font=FONT_BOX,
    )
    draw.multiline_text(
        (center_x, y0 + 70),
        body,
        anchor="ma",
        align="center",
        spacing=6,
        fill=hex_to_rgb(COLORS["muted"]),
        font=FONT_SMALL,
    )


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], dashed: bool = False) -> None:
    color = hex_to_rgb(COLORS["ink"])
    if dashed:
        dash = 16
        gap = 10
        if start[0] == end[0]:
            y = start[1]
            direction = 1 if end[1] > start[1] else -1
            while (direction == 1 and y < end[1]) or (direction == -1 and y > end[1]):
                next_y = y + direction * min(dash, abs(end[1] - y))
                draw.line((start[0], y, start[0], next_y), fill=color, width=4)
                y = next_y + direction * gap
        else:
            x = start[0]
            direction = 1 if end[0] > start[0] else -1
            while (direction == 1 and x < end[0]) or (direction == -1 and x > end[0]):
                next_x = x + direction * min(dash, abs(end[0] - x))
                draw.line((x, start[1], next_x, start[1]), fill=color, width=4)
                x = next_x + direction * gap
    else:
        draw.line((*start, *end), fill=color, width=4)

    arrow_size = 14
    ex, ey = end
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        if end[0] >= start[0]:
            points = [(ex, ey), (ex - arrow_size, ey - 8), (ex - arrow_size, ey + 8)]
        else:
            points = [(ex, ey), (ex + arrow_size, ey - 8), (ex + arrow_size, ey + 8)]
    else:
        if end[1] >= start[1]:
            points = [(ex, ey), (ex - 8, ey - arrow_size), (ex + 8, ey - arrow_size)]
        else:
            points = [(ex, ey), (ex - 8, ey + arrow_size), (ex + 8, ey + arrow_size)]
    draw.polygon(points, fill=color)


def draw_polyline_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    dashed: bool = False,
) -> None:
    if len(points) < 2:
        return

    color = hex_to_rgb(COLORS["ink"])
    if dashed:
        for start, end in zip(points, points[1:]):
            draw_arrow(draw, start, end, dashed=True)
        return

    draw.line(points, fill=color, width=4)
    start = points[-2]
    end = points[-1]
    arrow_size = 14
    ex, ey = end
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        if end[0] >= start[0]:
            arrow_points = [(ex, ey), (ex - arrow_size, ey - 8), (ex - arrow_size, ey + 8)]
        else:
            arrow_points = [(ex, ey), (ex + arrow_size, ey - 8), (ex + arrow_size, ey + 8)]
    else:
        if end[1] >= start[1]:
            arrow_points = [(ex, ey), (ex - 8, ey - arrow_size), (ex + 8, ey - arrow_size)]
        else:
            arrow_points = [(ex, ey), (ex - 8, ey + arrow_size), (ex + 8, ey + arrow_size)]
    draw.polygon(arrow_points, fill=color)


def generate_diagram() -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGBA", (2200, 1400), hex_to_rgb(COLORS["bg"]))
    draw = ImageDraw.Draw(image)

    draw.text(
        (110, 70),
        "Intent Compression Architecture (ICA)",
        fill=hex_to_rgb(COLORS["ink"]),
        font=FONT_TITLE,
    )
    draw.text(
        (112, 128),
        "Clarification-first control layer for reliable LLM systems",
        fill=hex_to_rgb(COLORS["muted"]),
        font=FONT_SUBTITLE,
    )

    top_y = 220
    top_boxes = {
        "user": (90, top_y, 420, top_y + 150),
        "intent": (500, top_y, 920, top_y + 150),
        "score": (1000, top_y, 1420, top_y + 150),
        "gate": (1500, top_y, 1970, top_y + 170),
    }

    draw_block(draw, top_boxes["user"], "User Query", "Raw request,\ncontext, constraints", "purple_fill", "purple_stroke")
    draw_block(
        draw,
        top_boxes["intent"],
        "Intent Hypothesis Generation",
        "Infer plausible meanings\nand answer deltas",
        "blue_fill",
        "blue_stroke",
    )
    draw_block(
        draw,
        top_boxes["score"],
        "Ambiguity + Risk Scoring",
        "Estimate uncertainty,\nsafety risk, and entropy",
        "blue_fill",
        "blue_stroke",
    )
    draw_block(
        draw,
        top_boxes["gate"],
        "Expected Utility Gate",
        "Choose route using\nexpected gain minus cost",
        "amber_fill",
        "amber_stroke",
    )

    draw_arrow(draw, (420, 295), (500, 295))
    draw_arrow(draw, (920, 295), (1000, 295))
    draw_arrow(draw, (1420, 295), (1500, 295))

    formula_box = (1540, 415, 1930, 520)
    draw.rounded_rectangle(
        formula_box,
        radius=18,
        fill=hex_to_rgb(COLORS["panel"]),
        outline=hex_to_rgb(COLORS["line"]),
        width=3,
    )
    draw.text(
        ((formula_box[0] + formula_box[2]) / 2, 442),
        "Ask iff  max U(q | x) > tau",
        anchor="ma",
        fill=hex_to_rgb(COLORS["ink"]),
        font=FONT_BODY,
    )
    draw.multiline_text(
        ((formula_box[0] + formula_box[2]) / 2, 478),
        "U(q | x) = E_r [L_direct - L_after] - C(q)",
        anchor="ma",
        align="center",
        spacing=5,
        fill=hex_to_rgb(COLORS["muted"]),
        font=FONT_SMALL,
    )

    route_y = 650
    route_boxes = {
        "answer": (120, route_y, 480, route_y + 150),
        "clarify": (560, route_y, 920, route_y + 150),
        "premise": (1000, route_y, 1360, route_y + 150),
        "refuse": (1440, route_y, 1800, route_y + 150),
    }

    draw_block(draw, route_boxes["answer"], "Answer Directly", "Low ambiguity,\nlow risk", "green_fill", "green_stroke")
    draw_block(draw, route_boxes["clarify"], "Ask Clarifier", "High ambiguity,\nlow risk", "purple_fill", "purple_stroke")
    draw_block(
        draw,
        route_boxes["premise"],
        "Premise-check Clarifier",
        "Clarify or strip false\ncausal framing",
        "purple_fill",
        "purple_stroke",
    )
    draw_block(
        draw,
        route_boxes["refuse"],
        "Refuse / Redirect",
        "High risk when\nclarification will not help",
        "red_fill",
        "red_stroke",
    )

    gate_center_x = (top_boxes["gate"][0] + top_boxes["gate"][2]) / 2
    for target in ("answer", "clarify", "premise", "refuse"):
        box = route_boxes[target]
        draw_arrow(draw, (int(gate_center_x), top_boxes["gate"][3]), ((box[0] + box[2]) // 2, box[1]))

    lower_y = 1020
    state_box = (690, lower_y, 1230, lower_y + 150)
    answer_box = (1310, lower_y, 1860, lower_y + 150)

    draw_block(
        draw,
        state_box,
        "Confirmed / Safe Intent State",
        "Narrowed meaning,\nconstraints, safe mode",
        "teal_fill",
        "teal_stroke",
    )
    draw_block(
        draw,
        answer_box,
        "Conditioned Answer Generation",
        "Final answer or safe\ncompletion",
        "green_fill",
        "green_stroke",
    )

    for source in ("clarify", "premise"):
        box = route_boxes[source]
        draw_arrow(draw, ((box[0] + box[2]) // 2, box[3]), ((state_box[0] + state_box[2]) // 2, state_box[1]))

    draw_arrow(draw, ((state_box[0] + state_box[2]) // 2, state_box[3] - 75), (answer_box[0], answer_box[1] + 75))
    draw_polyline_arrow(
        draw,
        [
            ((route_boxes["answer"][0] + route_boxes["answer"][2]) // 2, route_boxes["answer"][3]),
            ((route_boxes["answer"][0] + route_boxes["answer"][2]) // 2, 950),
            (answer_box[0] - 40, 950),
            (answer_box[0], answer_box[1] + 75),
        ],
    )
    draw_polyline_arrow(
        draw,
        [
            ((route_boxes["refuse"][0] + route_boxes["refuse"][2]) // 2, route_boxes["refuse"][3]),
            ((route_boxes["refuse"][0] + route_boxes["refuse"][2]) // 2, 930),
            (answer_box[2] + 30, 930),
            (answer_box[2], answer_box[1] + 75),
        ],
    )

    legend_x, legend_y = 70, 1080
    draw.text((legend_x, legend_y - 36), "Legend", fill=hex_to_rgb(COLORS["ink"]), font=FONT_BOX)
    legend_items = [
        ("blue_fill", "State estimation"),
        ("amber_fill", "Decision gate"),
        ("purple_fill", "User interaction"),
        ("teal_fill", "Narrowed intent state"),
        ("green_fill", "Answer path"),
        ("red_fill", "Safety redirect / refusal"),
    ]
    for idx, (fill_key, label) in enumerate(legend_items):
        y = legend_y + idx * 38
        draw.rounded_rectangle(
            (legend_x, y, legend_x + 24, y + 24),
            radius=6,
            fill=hex_to_rgb(COLORS[fill_key]),
            outline=hex_to_rgb(COLORS["line"]),
            width=2,
        )
        draw.text((legend_x + 40, y + 2), label, fill=hex_to_rgb(COLORS["muted"]), font=FONT_SMALL)

    footer_note = (
        "Routing principle: ask only when the highest-utility clarifier clears the threshold; "
        "otherwise answer directly, premise-check, or refuse/redirect."
    )
    draw.text((70, 1325), footer_note, fill=hex_to_rgb(COLORS["muted"]), font=FONT_SMALL)

    output = DIAGRAM_DIR / "architecture.png"
    image.convert("RGB").save(output)
    return output


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph_style(document: Document, name: str, size: int, bold: bool = False, color: str = "1F2937"):
    if name in document.styles:
        style = document.styles[name]
    else:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    return style


def build_docx(diagram_path: Path) -> Path:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    add_paragraph_style(document, "ICA Title", 22, bold=True, color="20354F")
    add_paragraph_style(document, "ICA Subtitle", 12, bold=False, color="5B6470")
    add_paragraph_style(document, "ICA Heading 1", 16, bold=True, color="20354F")
    add_paragraph_style(document, "ICA Heading 2", 13, bold=True, color="20354F")
    add_paragraph_style(document, "ICA Callout", 10, bold=False, color="20354F")

    header = section.header.paragraphs[0]
    header.text = "Intent Compression Architecture | Engineering Design Proposal"
    header.style = document.styles["Normal"]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header.runs:
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor.from_string("5B6470")

    footer = section.footer.paragraphs[0]
    footer.text = "Paul Maddison | Clarification-first control layer for reliable LLM systems"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if footer.runs:
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor.from_string("5B6470")

    p = document.add_paragraph(style="ICA Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Intent Compression Architecture").bold = True

    p = document.add_paragraph(style="ICA Subtitle")
    p.add_run("A Clarification-First Control Layer for Reliable LLM Systems")
    p = document.add_paragraph(style="ICA Subtitle")
    p.add_run("Engineering design proposal")

    meta = document.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    meta.columns[0].width = Inches(1.9)
    meta.columns[1].width = Inches(4.8)
    meta_rows = [
        ("Author", "Paul Maddison"),
        ("Short name", "ICA"),
        ("Canonical artifacts", "README.md, architecture diagram, proposal DOCX/PDF, evaluation scaffold"),
    ]
    for row, values in zip(meta.rows, meta_rows):
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(10)
            if idx == 0:
                set_cell_shading(cell, "E8F1FB")
                cell.paragraphs[0].runs[0].bold = True

    document.add_paragraph("")

    callout = document.add_paragraph(style="ICA Callout")
    callout.alignment = WD_ALIGN_PARAGRAPH.LEFT
    callout.add_run(
        "Executive summary. ICA is a pre-generation control layer that decides whether clarification is worth the cost "
        "before the model commits to an answer. It treats ambiguity as uncertainty over latent user intent, estimates "
        "expected utility over possible clarification replies, and routes the request to direct answer, clarifier, "
        "premise-check, or refusal/redirect as appropriate."
    )

    def heading(text: str, level: int = 1) -> None:
        style = "ICA Heading 1" if level == 1 else "ICA Heading 2"
        h = document.add_paragraph(style=style)
        h.add_run(text)

    heading("1. Architecture overview")
    for text in [
        "A common failure mode in LLM systems is unresolved ambiguity. Users often ask questions that admit multiple plausible interpretations, while the system is optimized to answer immediately. The result is a broad first answer, a user correction, and an unnecessary second pass.",
        "ICA inserts a control layer between user input and final generation. The layer estimates likely intent hypotheses, scores ambiguity and risk, and decides whether clarification improves the expected outcome enough to justify the extra turn.",
    ]:
        document.add_paragraph(text)

    document.add_picture(str(diagram_path), width=Inches(6.55))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Figure 1. ICA control flow with expected-utility routing and embedded legend.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("5B6470")

    heading("2. Core decision policy")
    document.add_paragraph(
        "Minimal rule: ask a clarifying question only when expected improvement in answer quality or safety exceeds the cost of another turn."
    )
    eq = document.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_run = eq.add_run("U(q | x) = E_r [L_direct(x) - L_after(x, q, r)] - C(q)")
    eq_run.italic = True
    eq_run.font.size = Pt(11)
    eq_run.font.name = "Consolas"
    eq2 = document.add_paragraph()
    eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq2_run = eq2.add_run("Ask iff max_q U(q | x) > tau_domain")
    eq2_run.italic = True
    eq2_run.font.size = Pt(11)
    eq2_run.font.name = "Consolas"
    document.add_paragraph(
        "This formulation makes the unknown user reply explicit by taking an expectation over possible clarification responses. The domain threshold allows the same architecture to be tuned differently for coding, medical, legal, finance, or customer-support use cases."
    )

    heading("3. Estimation and calibration")
    document.add_paragraph(
        "The schema fields intent_entropy_bits and intent_hypotheses[].probability are not meant to be accepted as magical model self-knowledge. In production they should be calibrated control-layer estimates: generate mutually exclusive answer-changing intent hypotheses, include an other bucket, estimate probabilities from a calibrated classifier or semantic clusters of repeated samples, and then compute entropy mechanically as -sum p log2 p."
    )
    document.add_paragraph(
        "Candidate utility should be treated the same way. The repository policy module uses provider-estimated clarifier benefit as one input, then applies deterministic local adjustments for token cost, latency cost, turn friction, and optional risk adjustment. That keeps the ask-vs-answer decision auditable even when upstream estimates are noisy."
    )
    document.add_paragraph(
        "Tau should be tuned, not guessed. The simplest calibration is a grid search on held-out labeled prompts that minimizes the combined loss from unnecessary clarification, false direct answers, false refusals, token cost, latency, abandonment, and safety risk. Operational dashboards should watch both failure modes: over-clarification when tau is too low and silent wrong-funnel answers when tau is too high."
    )

    heading("4. Compression interpretation")
    document.add_paragraph(
        "ICA treats ambiguity as entropy over possible user intents. Before clarification, the system reasons over H(I | x). After clarification q and user reply r, the system reasons over H(I | x, q, r). The expected value of clarification is the reduction in intent entropy, tempered by interaction cost."
    )
    document.add_paragraph(
        "This is why the term intent compression is mathematically defensible: the controller narrows the intent distribution before generation rather than letting the answering model hedge across multiple meanings."
    )

    heading("5. Routing policy")
    table = document.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Ambiguity", "Risk", "Default action", "Operational note"]
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = label
        set_cell_shading(cell, "E8F1FB")
        cell.paragraphs[0].runs[0].bold = True
    rows = [
        ("Low", "Low", "Answer directly", "No clarification needed."),
        ("High", "Low", "Ask one high-value clarifier", "Clarify only if the best question clears the threshold."),
        ("Low", "High", "Direct safe completion, premise-check, or refuse/redirect", "Do not ask the user to reconfirm harmful intent if the safe response would not change."),
        ("High", "High", "Strict clarification, constrained answer, or refuse/redirect", "High uncertainty plus high downside requires tighter control."),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_values):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.add_paragraph(
        "A practical implication is that clarification is not always the right response to risk. Sometimes the cleanest and safest route is a direct refusal with a helpful redirect."
    )

    heading("6. Prior-art positioning")
    for text in [
        "ICA is closest to selective-clarification and ambiguous-question-answering work such as CLAM, CLAMBER, and AmbigQA. It also relates to ReAct-style agent loops and uncertainty methods such as self-consistency and semantic entropy.",
        "The contribution claimed here is narrower and more architectural: clarification is specified as a pre-generation routing contract with explicit ambiguity, risk, utility, threshold, route, and counter-metric fields. The novelty is not asking questions; it is engineering ask-vs-answer as a calibrated control layer.",
        "Primary references: CLAM (arxiv.org/abs/2212.07769), CLAMBER (arxiv.org/abs/2405.12063), AmbigQA (arxiv.org/abs/2004.10645), ReAct (arxiv.org/abs/2210.03629), Self-Consistency (arxiv.org/abs/2203.11171), and Semantic Uncertainty (arxiv.org/abs/2302.09664).",
    ]:
        document.add_paragraph(text)

    document.add_page_break()
    heading("7. Implementation contract")
    document.add_paragraph(
        "The control layer should emit a structured decision object rather than a free-form paragraph. This proposal includes a JSON schema and reference example in spec/clarifier_output.schema.json and spec/clarifier_output.example.json."
    )
    document.add_paragraph(
        "The repository also includes a lightweight schema validation path so the example can be tested against the contract rather than presented as a static mockup. Local validation is the canonical reproducibility path."
    )
    code = document.add_paragraph()
    code.style = document.styles["Normal"]
    for line in [
        "{",
        '  "ambiguity_score": 0.74,',
        '  "risk_score": 0.22,',
        '  "decision": "ask_clarifier",',
        '  "clarifying_question": "When you say propaganda, do you mean persuasive political advocacy, coordinated deceptive messaging, or something else?",',
        '  "answer_constraints": ["avoid loaded framing", "define terms before conclusion"]',
        "}",
    ]:
        run = code.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)

    heading("8. Hidden baseline cost")
    document.add_paragraph(
        "A short first answer is not necessarily an efficient answer. In ambiguous prompts, a baseline system can choose the wrong interpretation, give a broad or misleading answer, and force the user into a correction funnel before the real issue is finally exposed."
    )
    document.add_paragraph(
        "This is not only a retry problem. It is a wrong-funnel problem. The user either spends turns arguing the model into discovering the ambiguous term, or leaves with a partially wrong answer without ever learning what caused the mismatch."
    )
    document.add_paragraph(
        "For that reason, ICA should be evaluated on tokens to resolved intent, definition-discovery turn, correction-funnel depth, and user correction burden, not just on the cost of the first assistant response."
    )

    heading("9. Evaluation package")
    document.add_paragraph(
        "The repository now includes a benchmark prompt set, an evaluation protocol, a sample reporting format, and a first-pass pilot benchmark. The evaluation compares not only direct one-shot answers, but also direct answers followed by the repair funnel when the first answer misses the intended meaning. The next credibility jump is a multi-rater or live-user benchmark."
    )
    document.add_paragraph(
        "The published pilot is designed to test ambiguous-prompt handling, not to estimate production-wide clarification frequency."
    )
    eval_table = document.add_table(rows=6, cols=2)
    eval_table.style = "Table Grid"
    eval_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    eval_rows = [
        ("Artifact", "Purpose"),
        ("examples/ambiguous_prompts.csv", "Starter prompt set covering low-risk ambiguity, high-risk ambiguity, and premise-risk cases."),
        ("eval/README.md", "Protocol, scoring rubric, success criteria, and counter-metrics."),
        ("eval/sample_results.md", "Template for benchmark reporting without fabricating results."),
        ("eval/pilot_results.md", "Human-readable first-pass pilot benchmark report."),
        ("eval/pilot_results.csv", "Machine-readable pilot results table for analysis and reuse."),
    ]
    for row_idx, values in enumerate(eval_rows):
        for col_idx, value in enumerate(values):
            cell = eval_table.rows[row_idx].cells[col_idx]
            cell.text = value
            if row_idx == 0:
                set_cell_shading(cell, "E8F1FB")
                cell.paragraphs[0].runs[0].bold = True

    document.add_paragraph(
        "Key primary metrics include first assistant-message tokens, total tokens to resolved intent, definition-discovery turn, retry count, correctness, clarity, and premise handling. Key counter-metrics include over-clarification rate, false direct-answer rate, silent-failure proxy, false refusal rate, and clarification bias."
    )
    document.add_paragraph(
        "The pilot utility proxy is transparent rather than implicit: quality = (correctness + clarity + safety) / 3; utility_proxy = quality - 0.01 * total_tokens - 0.5 * retries. This proxy rewards judged final-answer quality and penalizes token cost and retry burden. It does not measure live satisfaction, abandonment, wall-clock latency, or revenue impact."
    )
    document.add_paragraph(
        "The current 25-prompt pilot is deliberately ambiguity-heavy and routes 20 of 25 prompts to ask_clarifier. That 80 percent clarification rate is defensible for a stress set, but it is not a desired production rate. On representative traffic, clarification rate, over-clarification rate, unnecessary clarification rate, and false direct-answer rate should become the primary tau-calibration signals."
    )
    document.add_paragraph(
        "The circularity risk is real. The current pilot is single-rater and uses evaluator-supplied clarification replies. The next benchmark should separate reply generation from scoring, use route-blind or independent review where feasible, report inter-rater agreement, and tune tau on a separate split from the headline evaluation."
    )

    heading("10. Strategic implication")
    document.add_paragraph(
        "At large scale, ICA is not only a reliability pattern. It becomes a clarification data flywheel. Each ambiguous query, clarifier, user reply, and final outcome produces a structured intent-resolution trace that is more valuable than an ordinary chat log because it captures what the user actually meant."
    )
    document.add_paragraph(
        "That data advantage only exists if the system logs and acts on the traces: capture candidate intents, probabilities, rejected clarifiers, tau, selected route, user reply, and final outcome; label resolved intent with privacy controls; recalibrate probabilities and thresholds per domain; retrain the policy on both good clarifiers and cases where the system should not have asked; and deploy changes through A/B tests that watch success metrics and counter-metrics."
    )
    document.add_paragraph(
        "Public OpenAI disclosures in 2026 describe ChatGPT as having more than 900 million weekly active users and more than 50 million consumer subscribers. At that scale, even rare ambiguity classes become common in absolute terms. A provider with that distribution can improve ambiguity detection, ask-vs-answer thresholds, neutral clarifier wording, risk handling, and future base-model behavior using real-world traces rather than only synthetic or expert-labeled data."
    )
    document.add_paragraph(
        "The architecture is copyable. The live, high-volume intent-resolution feedback loop is much harder to copy. That is why ICA can be understood not only as a UX improvement, but as a route by which market share becomes model-quality advantage."
    )

    heading("11. Adversarial and deployment guidance")
    document.add_paragraph(
        "ICA is most useful when deployed as a policy layer around an existing model stack. The orchestration logic should be explicit and testable, while uncertainty is isolated to model outputs and external-system calls such as retrieval, APIs, or mutable external state."
    )
    document.add_paragraph(
        "The controller should log ambiguity score, risk score, candidate clarifiers, expected utility estimates, final route, and user reply when applicable. Those traces are the data needed to improve the control layer over time."
    )
    document.add_paragraph(
        "A clarification layer is also a control surface. It can be probed by users who try to trigger or suppress clarification, launder harmful intent through ambiguity, steer unsafe options through clarifier replies, inject desired routing labels, overload the controller with many plausible intents, or poison future traces. Mitigations include internal thresholds, independent risk scoring, post-reply safety re-scoring, policy-controlled hypothesis generation, capped hypothesis sets with an other bucket, and separation between online traces and trusted training labels."
    )

    heading("12. Conclusion")
    document.add_paragraph(
        "ICA is a credible architecture pattern because it defines a control-layer problem that engineers can implement: infer intent hypotheses, quantify ambiguity, route by expected utility, narrow intent before generation when doing so is worth the cost, and build systems that are more precise, more reliable, easier to defend, and often cheaper to operate once wrong-funnel conversations are counted properly."
    )

    heading("Appendix A. Starter benchmark pack")
    document.add_paragraph(
        "The repository includes a lightweight evaluation package so the proposal can move from theory to early evidence without requiring a large research program."
    )
    appendix = document.add_table(rows=6, cols=2)
    appendix.style = "Table Grid"
    appendix.alignment = WD_TABLE_ALIGNMENT.LEFT
    appendix_rows = [
        ("Artifact", "Use"),
        ("examples/ambiguous_prompts.csv", "Seed prompt set covering coding, planning, legal, medical, finance, public reasoning, and safety-sensitive ambiguity."),
        ("eval/README.md", "Evaluation procedure, rating rubric, success criteria, and counter-metrics."),
        ("eval/sample_results.md", "Reporting format for summary metrics and per-prompt comparisons."),
        ("eval/pilot_results.md", "Published first-pass pilot benchmark report."),
        ("eval/pilot_results.csv", "Structured pilot dataset for re-analysis or extension."),
    ]
    for row_idx, values in enumerate(appendix_rows):
        for col_idx, value in enumerate(values):
            cell = appendix.rows[row_idx].cells[col_idx]
            cell.text = value
            if row_idx == 0:
                set_cell_shading(cell, "E8F1FB")
                cell.paragraphs[0].runs[0].bold = True

    heading("Pilot signal snapshot", level=2)
    document.add_paragraph(
        "In the current 25-prompt pilot, direct first-pass answers required a repair funnel in 19 cases. On that same prompt set, ICA reduced mean definition-discovery turn from 2.6 to 1.0 and reduced mean total tokens to satisfactory resolution from 91.68 on the repaired baseline path to 78.04."
    )
    document.add_paragraph(
        "That is the narrower but stronger efficiency claim: early clarification is not always shorter than a one-shot answer, but it is often cheaper than discovering the same ambiguity after the system has already committed to the wrong semantic funnel."
    )
    document.add_paragraph(
        "In repaired-baseline scoring, final answer quality is equalized with ICA only when the baseline needed repair. The repaired utility score still penalizes extra repair tokens and retry burden so delayed clarification does not receive a free tie."
    )
    document.add_paragraph(
        "Reproduction path: install dependencies from requirements.txt, or requirements.lock for a pinned replay, run the local validation helper script, inspect the generated CSV and Markdown report, then regenerate the proposal artifacts if needed."
    )
    document.add_paragraph(
        "Recommended next move: extend the single-rater pilot into a multi-rater or API-instrumented benchmark, then update the controller threshold and routing rules based on observed over-clarification, correction-funnel depth, silent-failure risk, false direct-answer, and false-refusal rates."
    )

    document.save(DOCX_PATH)
    return DOCX_PATH


def main() -> None:
    diagram = generate_diagram()
    build_docx(diagram)


if __name__ == "__main__":
    main()
